"""
Tests for the document ingestion pipeline.

Covers: loading, chunking, metadata, and error handling.
"""

import pytest
from pathlib import Path
from docx import Document as DocxDocument

from backend.core.ingestion import ingest, load_document, split_documents


# ─────────────────────────────────────────────
# Fixtures
# ─────────────────────────────────────────────

@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """
    Creates a temporary .docx file with known content for testing.
    tmp_path is a pytest built-in fixture that provides a
    temporary directory unique to each test run.
    """
    doc = DocxDocument()
    doc.add_heading("DocMind Test Document", level=1)
    doc.add_paragraph(
        "This is the first paragraph of the test document. "
        "It contains enough text to be meaningful for chunking tests. "
        "DocMind is a multi-agent document intelligence platform."
    )
    doc.add_paragraph(
        "This is the second paragraph. It discusses retrieval-augmented generation, "
        "also known as RAG. RAG combines semantic search with large language models "
        "to produce grounded, citation-backed answers from document collections."
    )
    doc.add_paragraph(
        "The third paragraph covers LangGraph, which is used for orchestrating "
        "multiple AI agents. Each agent specialises in one task: retrieval, "
        "summarisation, comparison, or evaluation."
    )

    file_path = tmp_path / "test_document.docx"
    doc.save(str(file_path))
    return file_path


# ─────────────────────────────────────────────
# Tests
# ─────────────────────────────────────────────

class TestLoadDocument:
    """Tests for the load_document function."""

    def test_loads_docx_successfully(self, sample_docx: Path):
        """A valid .docx file should load without errors."""
        documents = load_document(sample_docx)
        assert len(documents) > 0

    def test_loaded_document_has_content(self, sample_docx: Path):
        """Loaded documents should contain non-empty text."""
        documents = load_document(sample_docx)
        combined_text = " ".join(doc.page_content for doc in documents)
        assert len(combined_text) > 0

    def test_loaded_document_has_source_metadata(self, sample_docx: Path):
        """Every loaded document should have a 'source' key in metadata."""
        documents = load_document(sample_docx)
        for doc in documents:
            assert "source" in doc.metadata
            assert doc.metadata["source"] == sample_docx.name

    def test_missing_file_raises_error(self, tmp_path: Path):
        """Loading a non-existent file should raise FileNotFoundError."""
        missing = tmp_path / "does_not_exist.docx"
        with pytest.raises(FileNotFoundError):
            load_document(missing)

    def test_unsupported_format_raises_error(self, tmp_path: Path):
        """Loading an unsupported file format should raise ValueError."""
        unsupported = tmp_path / "file.xyz"
        unsupported.write_text("some content")
        with pytest.raises(ValueError):
            load_document(unsupported)


class TestSplitDocuments:
    """Tests for the split_documents function."""

    def test_returns_chunks(self, sample_docx: Path):
        """Splitting a document should return at least one chunk."""
        documents = load_document(sample_docx)
        chunks = split_documents(documents)
        assert len(chunks) > 0

    def test_chunks_have_content(self, sample_docx: Path):
        """Every chunk should have non-empty page_content."""
        documents = load_document(sample_docx)
        chunks = split_documents(documents)
        for chunk in chunks:
            assert len(chunk.page_content.strip()) > 0

    def test_chunk_size_respected(self, sample_docx: Path):
        """No chunk should exceed the configured chunk_size significantly."""
        from backend.config import settings
        documents = load_document(sample_docx)
        chunks = split_documents(documents)
        for chunk in chunks:
            # Allow slight overflow due to splitter boundary logic
            assert len(chunk.page_content) <= settings.chunk_size * 1.2


class TestIngest:
    """Tests for the end-to-end ingest function."""

    def test_ingest_returns_chunks(self, sample_docx: Path):
        """ingest() should return a non-empty list of chunks."""
        chunks = ingest(sample_docx)
        assert isinstance(chunks, list)
        assert len(chunks) > 0

    def test_ingest_chunks_have_metadata(self, sample_docx: Path):
        """All chunks from ingest() should have source metadata."""
        chunks = ingest(sample_docx)
        for chunk in chunks:
            assert "source" in chunk.metadata

    def test_ingest_missing_file_raises_error(self, tmp_path: Path):
        """ingest() on a missing file should raise FileNotFoundError."""
        missing = tmp_path / "missing.docx"
        with pytest.raises(FileNotFoundError):
            ingest(missing)
